#!/usr/bin/env python3
"""parse-docx.py — 把 Word 需求文档解析为 MindDeck 项目骨架

用法：
    python3 scripts/parse-docx.py <input.docx> <output-project-dir> [--id <project-id>]

例：
    python3 scripts/parse-docx.py spec.docx Projects/my-project --id my-project

依赖（一次性）：
    pip3 install --user python-docx

行为：
    1. 解析 .docx 内 H1/H2/H3 层级建树
    2. 第一个 H1（如「一、功能简介」）整段并进根节点 note
    3. 第二个 H1（如「二、需求描述」）下的 H2 → depth=1 父节点，H3 → depth=2 页面节点
    4. 每个节点的 description 字段聚合其下所有段落 + 列表项（保留 `• ` 前缀）
    5. 节点首张内嵌图 → image 字段；额外图标记为 [图N] 留在 description 末尾
    6. 节点内 <w:tbl> → tables 字段（每表 {headers, rows}）
    7. 图片以内容 sha1 命名导出到 <output>/screenshots/
    8. data.js 输出扩展 schema（含 description / tables / nav_targets）
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
import urllib.request
import urllib.error

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("❌ 缺少 python-docx，跑：pip3 install --user python-docx", file=sys.stderr)
    sys.exit(1)


# ----------------- 标题清洗 -----------------
TITLE_PREFIX_PATTERNS = [
    re.compile(r'^[一二三四五六七八九十百千]+、\s*'),     # 一、二、
    re.compile(r'^[IVXLCM]+\.\s*', re.IGNORECASE),        # I. II. iii.
    re.compile(r'^\d+[、\.]\s*'),                         # 1、 1.
    re.compile(r'^[（(]\d+[）)]\s*'),                     # (1) （1）
]


def clean_title(raw):
    t = (raw or '').strip()
    for p in TITLE_PREFIX_PATTERNS:
        t = p.sub('', t, count=1)
    return t.replace('\n', ' ').strip()


# ----------------- 段落 / 块解析 -----------------
def get_para_style(p):
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        return ''
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        return ''
    return ps.get(qn('w:val')) or ''


def get_para_text(p):
    return ''.join(t.text or '' for t in p.iter(qn('w:t')))


def get_inline_image_rids(p):
    """返回该段落内所有 inline image 的 rId 列表（按出现顺序，可能多张）"""
    rids = []
    for blip in p.iter(qn('a:blip')):
        rid = blip.get(qn('r:embed'))
        if rid:
            rids.append(rid)
    return rids


def parse_table(tbl):
    """返回 {headers, rows}；headers 取第一行（若像表头），否则全归 rows"""
    rows = []
    for tr in tbl.findall(qn('w:tr')):
        cells = []
        for tc in tr.findall(qn('w:tc')):
            text = '\n'.join(
                ''.join(t.text or '' for t in p.iter(qn('w:t')))
                for p in tc.findall(qn('w:p'))
            ).strip()
            cells.append(text)
        rows.append(cells)
    if not rows:
        return None
    return {'headers': rows[0], 'rows': rows[1:]}


# ----------------- 图片导出 -----------------
def export_image(part, screenshots_dir, exported):
    """part: docx ImagePart；返回保存的文件名"""
    blob = part.blob
    digest = hashlib.sha1(blob).hexdigest()
    if digest in exported:
        return exported[digest]
    ext = (part.partname.rsplit('.', 1)[-1] or 'png').lower()
    if ext not in ('png', 'jpg', 'jpeg', 'gif', 'webp'):
        ext = 'png'
    fname = f'{digest}.{ext}'
    fpath = os.path.join(screenshots_dir, fname)
    if not os.path.exists(fpath):
        with open(fpath, 'wb') as f:
            f.write(blob)
    exported[digest] = fname
    return fname


# ----------------- 节点构造 -----------------
class Node:
    __slots__ = ('uid', 'depth', 'title', 'rawTitle', 'note', 'image',
                 'description_parts', 'extra_images', 'tables', 'children',
                 'nav_targets')

    def __init__(self, depth, raw_title):
        self.uid = None  # 后填
        self.depth = depth
        self.rawTitle = raw_title
        self.title = clean_title(raw_title)
        self.note = None
        self.image = None
        self.description_parts = []  # list of strings (含列表项前缀)
        self.extra_images = []        # 额外图片文件名（首张归 image，其余在此）
        self.tables = []
        self.children = []
        self.nav_targets = None

    def add_para(self, text, list_style=False):
        if not text or not text.strip():
            return
        if list_style:
            text = '• ' + text.strip()
        self.description_parts.append(text)

    def add_image(self, fname):
        if self.image is None:
            self.image = fname
        else:
            self.extra_images.append(fname)

    def add_table(self, t):
        if t:
            self.tables.append(t)

    def to_dict(self, with_children=True):
        desc_parts = list(self.description_parts)
        if self.extra_images:
            desc_parts.append('（另附 {} 张参考图）'.format(len(self.extra_images)))
        description = '\n\n'.join(desc_parts) if desc_parts else None
        out = {
            'uid': self.uid,
            'depth': self.depth,
            'title': self.title,
            'rawTitle': self.rawTitle,
            'note': self.note,
            'image': self.image,
            'description': description,
            'tables': self.tables if self.tables else None,
            'nav_targets': self.nav_targets,
            'children': [c.to_dict() for c in self.children] if with_children else [],
        }
        return out


# ----------------- 主解析 -----------------
def parse_docx(input_path, project_id, output_dir):
    doc = Document(input_path)
    screenshots_dir = os.path.join(output_dir, 'screenshots')
    os.makedirs(screenshots_dir, exist_ok=True)
    exported = {}  # sha1 → filename

    # 根节点：标题取 docx core_properties.title；没设就用文件名
    doc_title = (doc.core_properties.title or '').strip()
    if not doc_title:
        stem = os.path.splitext(os.path.basename(input_path))[0]
        # 去除尾部时间戳类后缀（- 加 8+ 位数字）
        stem = re.sub(r'-\d{8,}$', '', stem)
        doc_title = stem
    root = Node(0, doc_title)

    # 状态：current_h1_phase ∈ {'intro', 'requirements', None}
    # intro = 第一个 H1（功能简介）— 内容并入 root.note
    # requirements = 第二个 H1（需求描述）— 创建 H2/H3 树
    h1_phase = None
    current_h2 = None
    current_h3 = None
    intro_paras = []  # 第一个 H1 阶段的内容（最后变成 root.note）
    intro_tables = []

    # 当前内容应该追加到的目标节点
    def current_target():
        if h1_phase == 'requirements':
            return current_h3 or current_h2
        return None  # intro 阶段单独处理

    # 遍历 body 子元素
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            style = get_para_style(child)
            text = get_para_text(child).strip()
            rids = get_inline_image_rids(child)

            if style == 'Heading1':
                if h1_phase is None:
                    # 首个 H1（如「一、功能简介」）→ 进入 intro 阶段，内容并入 root.note
                    h1_phase = 'intro'
                else:
                    # 第二个 H1（如「二、需求描述」）→ 进入 requirements 阶段
                    h1_phase = 'requirements'
                    current_h2 = None
                    current_h3 = None
                continue

            if style == 'Heading2' and h1_phase == 'requirements':
                current_h2 = Node(1, text)
                current_h3 = None
                root.children.append(current_h2)
                continue

            if style == 'Heading3' and h1_phase == 'requirements':
                if current_h2 is None:
                    # H3 没爹 → 兜底建一个无名 H2
                    current_h2 = Node(1, '')
                    current_h2.title = '(未分组)'
                    root.children.append(current_h2)
                current_h3 = Node(2, text)
                current_h2.children.append(current_h3)
                continue

            # 非 heading 段落
            if h1_phase == 'intro':
                if text:
                    intro_paras.append(text)
                # intro 阶段的图也提取（万一有）
                for rid in rids:
                    try:
                        part = doc.part.related_parts[rid]
                        fname = export_image(part, screenshots_dir, exported)
                        if root.image is None:
                            root.image = fname
                    except KeyError:
                        pass

            elif h1_phase == 'requirements':
                tgt = current_target()
                if tgt is None:
                    continue  # H1=requirements 但还没遇到 H2，丢弃
                # 文本
                is_list = (style == 'ListParagraph')
                if text:
                    tgt.add_para(text, list_style=is_list)
                # 图片
                for rid in rids:
                    try:
                        part = doc.part.related_parts[rid]
                        fname = export_image(part, screenshots_dir, exported)
                        tgt.add_image(fname)
                    except KeyError:
                        pass

        elif tag == 'tbl':
            t = parse_table(child)
            if h1_phase == 'intro':
                if t:
                    intro_tables.append(t)
            elif h1_phase == 'requirements':
                tgt = current_target()
                if tgt and t:
                    tgt.add_table(t)

    # intro 阶段：纯文本段落 → root.note；表格 → root.tables（避免重复）
    root.note = '\n\n'.join(intro_paras) if intro_paras else None
    if intro_tables:
        root.tables = intro_tables

    # pre-order 分配 uid
    assign_uids(root)

    # 项目元信息
    meta = {'id': project_id, 'title': root.title or project_id, 'sub': None}
    return root, meta, exported


# ----------------- LLM 重构 + nav_targets 抽取 -----------------
ZHIPU_CHAT_URL = 'https://open.bigmodel.cn/api/paas/v4/chat/completions'


def call_zhipu(prompt, api_key, model='glm-4-flash', max_retry=3):
    """同步调智谱 chat completion，含简单退避重试。失败抛异常。"""
    body = json.dumps({
        'model': model,
        'messages': [{'role': 'user', 'content': prompt}],
    }).encode('utf-8')
    last_err = None
    for attempt in range(max_retry + 1):
        req = urllib.request.Request(
            ZHIPU_CHAT_URL,
            data=body,
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {api_key}',
            },
        )
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read())
                return data['choices'][0]['message']['content'] or ''
        except urllib.error.HTTPError as e:
            last_err = e
            if e.code == 429 and attempt < max_retry:
                wait = 2 ** attempt
                print(f'    429 限流，等 {wait}s 重试...')
                time.sleep(wait)
                continue
            raise
        except Exception as e:
            last_err = e
            if attempt < max_retry:
                time.sleep(1)
                continue
            raise
    raise last_err


def extract_json_array(text):
    """从 LLM 返回里抽 JSON 数组，容忍 ```json 包裹"""
    s = (text or '').strip()
    m = re.search(r'\[[\s\S]*\]', s)
    if not m:
        return []
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError:
        return []


def assign_uids(root):
    counter = [0]

    def walk(n):
        n.uid = 'n' + str(counter[0])
        counter[0] += 1
        for c in n.children:
            walk(c)
    walk(root)


def update_depths(node, depth):
    node.depth = depth
    for c in node.children:
        update_depths(c, depth + 1)


def collect_h3s(node, out):
    """收集所有 H3（depth=2 但可能因重组移动了，按是 H2 之下的 page 节点定义）"""
    if node.depth >= 2 and node.image:  # H3 通常有截图，作为 leaf-ish page
        out.append(node)
    for c in node.children:
        collect_h3s(c, out)


def llm_restructure(root, entry_title, api_key):
    """显式入口 + LLM 抽取「H3 描述里提到的跳转目标 H2」→ re-parent。
    返回 (re-parented count, nav_targets dict {h3_uid: [{label, trigger}]})
    """
    # 找入口 H2
    entry = next(
        (c for c in root.children if c.title == entry_title or entry_title in c.title or c.title in entry_title),
        None,
    )
    if not entry:
        print(f'⚠️  --entry "{entry_title}" 在 H2 中未找到（候选: {[c.title for c in root.children]}），跳过重构')
        return 0, {}

    # 候选目标 = 除入口外的所有 H2
    h2_others = [c for c in root.children if c is not entry]
    if not h2_others:
        return 0, {}

    # 收集所有 H3（带 description 的页面节点）
    all_h3 = []
    for h2 in [entry] + h2_others:
        for h3 in h2.children:
            if h3.description_parts:
                all_h3.append(h3)

    if not all_h3:
        print('⚠️  没有带描述的 H3，无法做语义匹配')
        return 0, {}

    print(f'  调用 glm-4-flash 分析 {len(all_h3)} 个 H3 → {len(h2_others)} 个候选目标...')

    def h3_summary(h):
        desc = ' '.join(h.description_parts)[:280]
        return f'[{h.uid}] "{h.title}" — {desc}'
    h3_block = '\n'.join(h3_summary(h) for h in all_h3)
    targets_block = '\n'.join(f'- "{t.title}"' for t in h2_others)

    prompt = f'''下面是产品需求文档解析出的页面节点。每个 H3 是一个具体页面，描述里可能包含跳转语句（如"点击金刚区进入公募基金"、"跳到 X 页"）。

H3 节点：
{h3_block}

候选跳转目标（待挂载的页面模块）：
{targets_block}

任务：识别哪些 H3 描述中**明确**提到跳转到候选目标。返回严格 JSON 数组：
[{{"h3_uid": "n10", "target": "公募基金", "trigger": "点击金刚区"}}, ...]

要求：
- 只返回**明确跳转语句**的匹配，模糊不返回
- target 必须**原样**用候选列表的字符串
- trigger 是触发动作（< 20 字），用于 UI 提示
- 同一目标如被多个 H3 引用，只返回第一个最直接的

只输出 JSON 数组，无其他说明文字。'''

    try:
        raw = call_zhipu(prompt, api_key)
    except Exception as e:
        print(f'⚠️  LLM 调用失败: {e}，跳过重构')
        return 0, {}

    pairs = extract_json_array(raw)
    print(f'  LLM 返回 {len(pairs)} 条匹配')

    # nav_targets 按 H3 uid 聚合
    nav_targets = {}
    re_parent_count = 0

    h3_by_uid = {h.uid: h for h in all_h3}
    h2_by_title = {h.title: h for h in h2_others}

    for p in pairs:
        h3_uid = p.get('h3_uid')
        target_title = p.get('target')
        trigger = p.get('trigger', '')
        h3 = h3_by_uid.get(h3_uid)
        target_node = h2_by_title.get(target_title)
        if not h3 or not target_node:
            continue
        # 记录 nav_target
        nav_targets.setdefault(h3_uid, []).append({
            'label': target_title,
            'trigger': trigger,
            'uid': target_node.uid,  # 暂存原 uid，re-parent 后 uid 不变
        })
        # 仅当 target 仍在 root.children 时才 re-parent（避免重复挂载）
        if target_node in root.children:
            root.children.remove(target_node)
            h3.children.append(target_node)
            re_parent_count += 1
            print(f'  ↳ "{target_node.title}" 挂到 "{h3.title}" (触发: {trigger})')

    # 把 nav_targets 写入对应 H3 节点
    for h3_uid, targets in nav_targets.items():
        h3 = h3_by_uid[h3_uid]
        h3.nav_targets = targets

    # 重新计算 depth + 重新分配 uid（pre-order）
    update_depths(root, 0)
    # uid 不重排（保持原 uid 让 nav_targets.uid 引用稳定）

    return re_parent_count, nav_targets


# ----------------- 输出 data.js -----------------
def render_data_js(meta, root):
    return (
        f'window.PROJECT_META = {json.dumps(meta, ensure_ascii=False)};\n'
        f'window.PROTOTYPE_TREE = {json.dumps(root.to_dict(), ensure_ascii=False, indent=2)};\n'
    )


# ----------------- CLI -----------------
def sanitize_id(name):
    base = os.path.splitext(os.path.basename(name))[0].lower()
    s = re.sub(r'[^a-z0-9一-鿿]+', '-', base)  # 保留中文
    return s.strip('-') or 'project'


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('input', help='.docx 输入文件')
    ap.add_argument('output', help='输出项目目录')
    ap.add_argument('--id', help='项目 ID（默认从文件名推断）')
    ap.add_argument('--entry', help='入口 H2 标题（启用 LLM 重构。例：--entry "APP首页"）')
    ap.add_argument('--no-llm', action='store_true', help='跳过 LLM 重构，仅按 doc 结构输出')
    args = ap.parse_args()

    if not os.path.exists(args.input):
        print(f'❌ 输入文件不存在: {args.input}', file=sys.stderr)
        sys.exit(1)

    project_id = args.id or sanitize_id(args.input)
    os.makedirs(args.output, exist_ok=True)

    print(f'解析中：{args.input}')
    root, meta, exported = parse_docx(args.input, project_id, args.output)

    # LLM 重构（可选）
    if args.entry and not args.no_llm:
        api_key = os.environ.get('ZHIPU_API_KEY')
        if not api_key:
            # 兜底从 ~/.claude/settings.json 读
            try:
                d = json.load(open(os.path.expanduser('~/.claude/settings.json')))
                api_key = (d.get('env') or {}).get('ANTHROPIC_AUTH_TOKEN', '')
            except Exception:
                pass
        if not api_key:
            print('⚠️  未找到 ZHIPU_API_KEY 环境变量，跳过 LLM 重构')
        else:
            print(f'LLM 重构（入口 = "{args.entry}"）...')
            n_reparented, navs = llm_restructure(root, args.entry, api_key)
            print(f'  重构: {n_reparented} 个 H2 已挂到 H3 下，{len(navs)} 个 H3 标记了 nav_targets')

    data_js_path = os.path.join(args.output, 'data.js')
    with open(data_js_path, 'w', encoding='utf-8') as f:
        f.write(render_data_js(meta, root))

    # 统计
    def count_nodes(n):
        return 1 + sum(count_nodes(c) for c in n.children)

    def count_with_image(n):
        c = 1 if n.image else 0
        return c + sum(count_with_image(ch) for ch in n.children)

    def count_with_desc(n):
        c = 1 if n.description_parts else 0
        return c + sum(count_with_desc(ch) for ch in n.children)

    print(f'✓ 解析完成')
    print(f'  项目 ID         : {project_id}')
    print(f'  根标题          : {root.title}')
    print(f'  节点总数        : {count_nodes(root)}')
    print(f'  H2 父节点       : {len(root.children)}')
    print(f'  含截图节点      : {count_with_image(root)}')
    print(f'  含描述节点      : {count_with_desc(root)}')
    print(f'  截图文件（去重）: {len(exported)}')
    print(f'  data.js         : {data_js_path}')
    print('')
    print('下一步：从 Template/ 复制 Prototype.html / app.js / locator.js 到此目录，浏览器打开。')


if __name__ == '__main__':
    main()
